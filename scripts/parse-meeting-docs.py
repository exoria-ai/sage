#!/usr/bin/env python3
"""
Parse ReGIS meeting documents (PDF and DOCX) into structured JSON.

Extracts meeting content organized by agenda items for search and retrieval.

Usage:
    source .venv/bin/activate
    python scripts/parse-meeting-docs.py
"""

import json
import re
import sys
from dataclasses import dataclass, asdict, field
from pathlib import Path
from typing import Optional

try:
    import fitz  # PyMuPDF for PDF
except ImportError:
    print("PyMuPDF not installed. Run: pip install pymupdf")
    sys.exit(1)

try:
    from docx import Document  # python-docx for DOCX
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


# Paths
SCRIPT_DIR = Path(__file__).parent
DATA_DIR = SCRIPT_DIR.parent / "data" / "meetings" / "regis"
PDF_DIR = DATA_DIR / "pdfs"
OUTPUT_FILE = DATA_DIR / "regis-meetings.json"


@dataclass
class MeetingChunk:
    """A chunk of meeting document content."""
    id: str
    meeting_id: str
    committee: str
    meeting_date: str
    document_type: str  # "agenda" or "minutes"
    item_number: str
    item_title: str
    chunk_type: str  # "header", "attendance", "announcement", "discussion", "action"
    text: str
    page_start: int = 1
    page_end: int = 1
    attendees: list[str] = field(default_factory=list)
    action_taken: Optional[str] = None


def extract_text_from_pdf(path: Path) -> tuple[str, int]:
    """Extract text from PDF file. Returns (text, page_count)."""
    doc = fitz.open(path)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts), len(text_parts)


def extract_text_from_docx(path: Path) -> tuple[str, int]:
    """Extract text from DOCX file. Returns (text, estimated_pages)."""
    doc = Document(path)
    text_parts = []
    for para in doc.paragraphs:
        text_parts.append(para.text)

    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells]
            text_parts.append("\t".join(row_text))

    full_text = "\n".join(text_parts)
    # Estimate pages (roughly 3000 chars per page)
    estimated_pages = max(1, len(full_text) // 3000 + 1)
    return full_text, estimated_pages


def extract_meeting_date(filename: str) -> Optional[str]:
    """Extract date from filename like 'regis-2024-01-17-minutes.pdf'."""
    match = re.search(r'(\d{4}-\d{2}-\d{2})', filename)
    return match.group(1) if match else None


def extract_document_type(filename: str) -> str:
    """Extract document type from filename."""
    if 'minutes' in filename.lower():
        return 'minutes'
    return 'agenda'


def extract_attendees(text: str) -> list[str]:
    """Extract attendee names from minutes text."""
    attendees = []

    # Look for "Members Present" section
    present_match = re.search(
        r'(?:Members?\s+Present|ReGIS\s+Members?\s+Present)[:\s]*[-–]?\s*(?:In\s+Alphabetical\s+Order)?[:\s]*\n(.*?)(?:\n\s*\d+\.\s|\n\s*[A-Z]{2,}|\Z)',
        text,
        re.IGNORECASE | re.DOTALL
    )

    if present_match:
        attendee_block = present_match.group(1)
        # Split by common delimiters
        names = re.split(r'[•\n,]', attendee_block)
        for name in names:
            name = name.strip()
            # Filter out non-names
            if name and len(name) > 3 and not name.startswith('Chair'):
                # Clean up organization suffixes
                name = re.sub(r',?\s*(City of|County|SID|FSSD|SCWA|STA|VFW|Travis AFB).*$', '', name, flags=re.IGNORECASE)
                name = name.strip()
                if name and len(name) > 3:
                    attendees.append(name)

    return attendees[:30]  # Cap at 30 attendees


def parse_agenda_items(text: str) -> list[dict]:
    """Parse agenda items from text."""
    items = []

    # Pattern for agenda items: "1." or "1)" or numbered items
    # Typical format:
    # 1. INTRODUCTIONS – FOR NEW ATTENDEES/GUESTS (5 minutes)
    # 2. APPROVAL OF THE OCTOBER 10, 2025, MEETING MINUTES (5 minutes)

    item_pattern = re.compile(
        r'^(\d+)\.\s+([A-Z][A-Z\s/\-–,&\']+?)(?:\s*[\(（].*?[\)）])?\s*$',
        re.MULTILINE
    )

    matches = list(item_pattern.finditer(text))

    for i, match in enumerate(matches):
        item_num = match.group(1)
        item_title = match.group(2).strip()

        # Get content between this item and next
        start_pos = match.end()
        end_pos = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        content = text[start_pos:end_pos].strip()

        # Handle sub-items (a, b, c or 4.a, 4.b)
        sub_items = []
        sub_pattern = re.compile(r'^([a-z])\.\s+(.+?)$', re.MULTILINE)
        for sub_match in sub_pattern.finditer(content):
            sub_items.append({
                'letter': sub_match.group(1),
                'text': sub_match.group(2).strip()
            })

        items.append({
            'number': item_num,
            'title': item_title,
            'content': content,
            'sub_items': sub_items
        })

    return items


def determine_chunk_type(item_title: str, content: str) -> str:
    """Determine the type of agenda item."""
    title_lower = item_title.lower()

    if 'introduction' in title_lower or 'attendee' in title_lower or 'guest' in title_lower:
        return 'attendance'
    elif 'approval' in title_lower and 'minutes' in title_lower:
        return 'approval'
    elif 'announcement' in title_lower or 'opening' in title_lower or 'job' in title_lower:
        return 'announcement'
    elif 'adjourn' in title_lower:
        return 'adjournment'
    elif 'round table' in title_lower or 'discussion' in title_lower:
        return 'discussion'
    elif 'update' in title_lower or 'report' in title_lower:
        return 'update'
    elif 'action' in title_lower or 'vote' in title_lower:
        return 'action'
    else:
        return 'general'


def parse_document(path: Path) -> list[MeetingChunk]:
    """Parse a meeting document into chunks."""
    chunks = []

    filename = path.name
    meeting_date = extract_meeting_date(filename)
    doc_type = extract_document_type(filename)

    if not meeting_date:
        print(f"  Warning: Could not extract date from {filename}")
        return chunks

    meeting_id = f"regis-{meeting_date}"

    # Extract text based on file type
    if path.suffix.lower() == '.pdf':
        text, page_count = extract_text_from_pdf(path)
    elif path.suffix.lower() == '.docx':
        text, page_count = extract_text_from_docx(path)
    else:
        print(f"  Skipping unsupported format: {path.suffix}")
        return chunks

    # Clean text
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Extract attendees (for minutes)
    attendees = []
    if doc_type == 'minutes':
        attendees = extract_attendees(text)

    # Parse agenda items
    items = parse_agenda_items(text)

    if not items:
        # If we couldn't parse structured items, create a single chunk
        chunk = MeetingChunk(
            id=f"{meeting_id}-{doc_type}-001",
            meeting_id=meeting_id,
            committee="ReGIS Members",
            meeting_date=meeting_date,
            document_type=doc_type,
            item_number="0",
            item_title="Full Document",
            chunk_type="general",
            text=text[:8000],  # Cap at 8000 chars
            page_start=1,
            page_end=page_count,
            attendees=attendees if doc_type == 'minutes' else [],
        )
        chunks.append(chunk)
    else:
        # Create chunk for each agenda item
        for i, item in enumerate(items):
            chunk_type = determine_chunk_type(item['title'], item['content'])

            # Build full text including sub-items
            full_text = item['content']
            if item['sub_items']:
                sub_text = "\n".join([f"  {s['letter']}. {s['text']}" for s in item['sub_items']])
                if sub_text not in full_text:
                    full_text = f"{full_text}\n{sub_text}"

            chunk = MeetingChunk(
                id=f"{meeting_id}-{doc_type}-{str(i+1).zfill(3)}",
                meeting_id=meeting_id,
                committee="ReGIS Members",
                meeting_date=meeting_date,
                document_type=doc_type,
                item_number=item['number'],
                item_title=item['title'],
                chunk_type=chunk_type,
                text=full_text[:4000],  # Cap each chunk
                page_start=1,
                page_end=page_count,
                attendees=attendees if i == 0 and doc_type == 'minutes' else [],
            )
            chunks.append(chunk)

    return chunks


def main():
    print("=" * 60)
    print("ReGIS Meeting Document Parser")
    print("=" * 60)

    if not PDF_DIR.exists():
        print(f"Error: PDF directory not found: {PDF_DIR}")
        print("Run scrape-regis-meetings.py first to download documents.")
        sys.exit(1)

    # Find all documents
    pdf_files = sorted(PDF_DIR.glob("*.pdf"))
    docx_files = sorted(PDF_DIR.glob("*.docx"))
    all_files = pdf_files + docx_files

    print(f"\nFound {len(all_files)} documents ({len(pdf_files)} PDF, {len(docx_files)} DOCX)")

    all_chunks = []
    meetings_seen = set()

    for path in all_files:
        print(f"\nProcessing: {path.name}")
        try:
            chunks = parse_document(path)
            all_chunks.extend(chunks)

            for chunk in chunks:
                meetings_seen.add(chunk.meeting_id)

            print(f"  Extracted {len(chunks)} chunks")
        except Exception as e:
            print(f"  Error: {e}")

    # Build output structure
    output = {
        "committee": "ReGIS Members",
        "description": "Solano Regional GIS Consortium member meetings",
        "source_url": "https://regis.solanocounty.com/regis-members-meetings/",
        "total_meetings": len(meetings_seen),
        "total_chunks": len(all_chunks),
        "date_range": {
            "earliest": min(c.meeting_date for c in all_chunks) if all_chunks else None,
            "latest": max(c.meeting_date for c in all_chunks) if all_chunks else None,
        },
        "chunks": [asdict(c) for c in all_chunks]
    }

    # Write output
    OUTPUT_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(OUTPUT_FILE, 'w') as f:
        json.dump(output, f, indent=2)

    print("\n" + "=" * 60)
    print("Parse Summary")
    print("=" * 60)
    print(f"Total meetings: {len(meetings_seen)}")
    print(f"Total chunks: {len(all_chunks)}")

    # Count by type
    by_doc_type = {}
    by_chunk_type = {}
    for chunk in all_chunks:
        by_doc_type[chunk.document_type] = by_doc_type.get(chunk.document_type, 0) + 1
        by_chunk_type[chunk.chunk_type] = by_chunk_type.get(chunk.chunk_type, 0) + 1

    print(f"\nBy document type:")
    for dt, count in sorted(by_doc_type.items()):
        print(f"  {dt}: {count}")

    print(f"\nBy chunk type:")
    for ct, count in sorted(by_chunk_type.items()):
        print(f"  {ct}: {count}")

    print(f"\nOutput: {OUTPUT_FILE}")
    print(f"Size: {OUTPUT_FILE.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
